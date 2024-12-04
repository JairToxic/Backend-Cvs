from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Inches
import os
import datetime
import requests
from io import BytesIO
from flask_cors import CORS

app = Flask(__name__)

# Habilitar CORS en toda la aplicación
CORS(app)

# Rutas de las carpetas
TEMPLATES_FOLDER = 'templates_doc'
GENERATED_FOLDER = 'generated_docs'
UPLOADS_FOLDER = 'uploads'

os.makedirs(TEMPLATES_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
os.makedirs(UPLOADS_FOLDER, exist_ok=True)

# Ruta del archivo de plantilla
TEMPLATE_FILE = os.path.join(TEMPLATES_FOLDER, 'plantilla.docx')

def replace_placeholders(doc, data):
    """
    Reemplaza placeholders {{clave}} en el documento Word, incluso en tablas separadas.
    """
    # Reemplazo en párrafos
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"  # Ejemplo: {{nombre}}
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Reemplazo en todas las tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"  # Ejemplo: {{nombre}}
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

def insert_image_from_url(doc, image_url):
    """Descarga la imagen desde una URL y la inserta en el lugar del marcador de imagen."""
    try:
        # Descargar la imagen desde la URL
        response = requests.get(image_url)
        if response.status_code == 200:
            # Crear un objeto de imagen en memoria
            image_stream = BytesIO(response.content)

            # Buscar el marcador de imagen y reemplazarlo con la imagen
            for paragraph in doc.paragraphs:
                if '{{INSERTAR_IMAGEN}}' in paragraph.text:
                    # Reemplaza el marcador con la imagen
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(2))  # Ajusta el tamaño de la imagen si es necesario
                    paragraph.text = paragraph.text.replace('{{INSERTAR_IMAGEN}}', '')  # Elimina el marcador
                    break  # Salir después de reemplazar la imagen
    except Exception as e:
        print(f"Error al descargar la imagen: {e}")

@app.route('/generate-doc', methods=['POST'])
def generate_document():
    try:
        # Verificar si existe la plantilla
        if not os.path.exists(TEMPLATE_FILE):
            return jsonify({"error": f"La plantilla {TEMPLATE_FILE} no existe"}), 404

        # Obtener los datos enviados en el POST
        data = request.json
        if not data:
            return jsonify({"error": "Los datos son obligatorios"}), 400

        # Cargar la plantilla
        doc = Document(TEMPLATE_FILE)

        # Reemplazar las variables en el documento
        replace_placeholders(doc, data)

        # Insertar la imagen si se encuentra en los datos
        image_url = data.get('imagen1')
        if image_url:
            insert_image_from_url(doc, image_url)

        # Generar un nombre único para el documento generado
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        output_file = os.path.join(GENERATED_FOLDER, f'documento_{timestamp}.docx')

        # Guardar el documento generado
        doc.save(output_file)

        # Enviar el archivo para su descarga inmediata
        return send_file(output_file, as_attachment=True, download_name=f'documento_{timestamp}.docx')

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
